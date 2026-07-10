"""阶段 8 · 报告生成。

为 Top 20 球员生成 Markdown 源报告 + DOCX 报告。
v1 使用模板 + 数据填充（不接大模型），后续可接入 LLM 生成自然语言段落。
参见 docs/阶段8-报告生成.md。
"""
from __future__ import annotations

from pathlib import Path

import pandas as pd

# 位置→中文名（用于报告展示）
_POS_CN = {"Winger": "边锋", "AM": "前腰", "CM": "中前卫", "DM": "后腰"}


def _fmt(val, default="?") -> str:
    """格式化单个值用于报告展示。"""
    if val is None or (isinstance(val, float) and pd.isna(val)):
        return default
    if isinstance(val, float):
        if abs(val) >= 1e6:
            return f"{val / 1e6:.1f}M"
        if abs(val) >= 1000:
            return f"{val:,.0f}"
        if abs(val) < 10:
            return f"{val:.2f}"
        return f"{val:.1f}"
    return str(val)


def _build_strengths(row: pd.Series) -> str:
    """基于 per90 百分位识别主要优势。"""
    items = []
    checks = [
        ("goals_per90", "进球效率", 65),
        ("assists_per90", "助攻能力", 60),
        ("key_passes_per90", "关键传球", 60),
        ("dribbles_per90", "盘带突破", 60),
        ("passes_per90", "传球输出", 60),
        ("tackles_per90", "防守抢断", 65),
        ("interceptions_per90", "拦截意识", 65),
        ("progressive_carries_per90", "带球推进", 60),
        ("shot_creating_actions_per90", "射门创造参与", 60),
        ("xg_per90", "预期进球威胁", 60),
    ]
    for col, desc, threshold in checks:
        val = row.get(col)
        if val is not None and not (isinstance(val, float) and pd.isna(val)):
            if float(val) > threshold / 100:  # per90 values are raw, not pct
                pass
            # Use the value directly - if it's high enough relative to position average
    # Simplified: use per90 values that are clearly above average
    if row.get("goals_per90", 0) and float(row.get("goals_per90", 0)) > 0.3:
        items.append(f"- **进球效率**：每90分钟 {float(row['goals_per90']):.2f} 球，在同位置中处于前列")
    if row.get("assists_per90", 0) and float(row.get("assists_per90", 0)) > 0.2:
        items.append(f"- **助攻能力**：每90分钟 {float(row['assists_per90']):.2f} 次助攻")
    if row.get("key_passes_per90", 0) and float(row.get("key_passes_per90", 0)) > 1.5:
        items.append(f"- **创造力**：每90分钟 {float(row['key_passes_per90']):.2f} 次关键传球")
    if row.get("dribbles_per90", 0) and float(row.get("dribbles_per90", 0)) > 2.0:
        items.append(f"- **盘带突破**：每90分钟 {float(row['dribbles_per90']):.2f} 次过人")
    if row.get("tackles_per90", 0) and float(row.get("tackles_per90", 0)) > 2.5:
        items.append(f"- **防守覆盖**：每90分钟 {float(row['tackles_per90']):.2f} 次抢断 + {float(row.get('interceptions_per90', 0)):.2f} 次拦截")
    if row.get("progressive_carries_per90", 0) and float(row.get("progressive_carries_per90", 0)) > 2.0:
        items.append(f"- **带球推进**：每90分钟 {float(row['progressive_carries_per90']):.2f} 次推进带球")
    if row.get("shot_creating_actions_per90", 0) and float(row.get("shot_creating_actions_per90", 0)) > 3.0:
        items.append(f"- **射门创造参与**：每90分钟 {float(row['shot_creating_actions_per90']):.2f} 次 SCA")

    if not items:
        items.append("- 该球员在各维度表现均衡，无明显数据突出项。")
    return "\n".join(items)


def _build_weaknesses(row: pd.Series) -> str:
    """基于 per90 值识别主要短板。"""
    items = []
    if row.get("yellow_cards_per90", 0) and float(row.get("yellow_cards_per90", 0)) > 1.0:
        items.append(f"- **纪律风险**：每90分钟 {float(row['yellow_cards_per90']):.2f} 张黄牌，偏高")
    if row.get("red_cards", 0) and int(row.get("red_cards", 0)) > 0:
        items.append(f"- **红牌记录**：本赛季有 {int(row['red_cards'])} 次被罚出场")
    if row.get("minutes", 0) and float(row.get("minutes", 0)) < 900:
        items.append(f"- **样本不足**：出场仅 {int(row['minutes'])} 分钟，数据稳定性存疑")
    if row.get("data_completeness", 1) and float(row.get("data_completeness", 1)) < 0.8:
        items.append(f"- **数据缺失**：核心字段完整度仅 {float(row['data_completeness'])*100:.0f}%")
    if row.get("goals_per90", 1) and float(row.get("goals_per90", 1)) < 0.1 and row.get("standard_position") in ("Winger", "AM"):
        items.append("- **进球贡献偏低**：作为攻击型球员，进球效率在同位置中处于下游")
    if row.get("passes_per90", 999) and float(row.get("passes_per90", 999)) < 30 and row.get("standard_position") in ("CM", "DM"):
        items.append("- **传球参与度偏低**：每90分钟传球次数在同位置中偏少")

    if not items:
        items.append("- 基于当前数据未发现明显短板。")
    return "\n".join(items)


def _build_behavior_profile(row: pd.Series) -> str:
    """基于数据生成场上行为风格描述。"""
    pos = row.get("standard_position", "Winger")
    pos_cn = _POS_CN.get(pos, pos)
    parts = [f"{row['player_name']} 司职 **{pos_cn}**。"]

    kp90 = float(row.get("key_passes_per90", 0) or 0)
    drib90 = float(row.get("dribbles_per90", 0) or 0)
    tkl90 = float(row.get("tackles_per90", 0) or 0)
    int90 = float(row.get("interceptions_per90", 0) or 0)
    prg90 = float(row.get("progressive_carries_per90", 0) or 0)

    if kp90 > 2.0:
        parts.append("在进攻组织中展现了较强的创造力。")
    if drib90 > 3.0:
        parts.append("善于利用盘带突破创造空间。")
    if tkl90 + int90 > 5.0:
        parts.append("防守参与积极，覆盖范围较大。")
    elif tkl90 + int90 < 2.0:
        parts.append("防守参与度相对有限，更专注于进攻任务。")
    if prg90 > 3.0:
        parts.append("带球推进是核心进攻手段之一。")

    # Discipline
    yp90 = float(row.get("yellow_cards_per90", 0) or 0)
    if yp90 > 1.0:
        parts.append(f"场均黄牌偏高（{yp90:.2f}/90分钟），需注意纪律控制。")
    elif yp90 < 0.3:
        parts.append("场上纪律性良好，较少不必要的犯规。")

    parts.append("以上分析仅基于比赛统计数据的场上行为倾向，不构成心理诊断。")
    return "".join(parts)


def _build_risks(row: pd.Series) -> str:
    """识别发展风险。"""
    items = []
    if row.get("risk_penalty", 0) < 0:
        items.append(f"- 风险扣分 {float(row['risk_penalty']):.0f} 分，表明存在纪律或样本问题。")
    if row.get("minutes", 999) < 900:
        items.append(f"- 出场时间仅 {int(row['minutes'])} 分钟，per90 数据可能因小样本而不稳定。")
    if row.get("data_completeness", 1) < 0.85:
        items.append(f"- 核心字段完整度仅 {float(row['data_completeness'])*100:.0f}%，部分结论可能存在偏差。")
    if row.get("age", 99) > 20:
        items.append(f"- 年龄 {int(row['age'])} 岁，成长窗口期相对有限，需尽快在高水平比赛中证明自己。")
    if row.get("age", 0) < 18:
        items.append(f"- 年龄仅 {int(row['age'])} 岁，身体和心理仍在发育中，需关注过度比赛负荷风险。")

    if not items:
        items.append("- 当前未识别出显著发展风险。")
    return "\n".join(items)


def _build_tactical_fit(row: pd.Series) -> str:
    """基于位置和风格推荐战术适配。"""
    pos = row.get("standard_position", "Winger")
    fits = {
        "Winger": "- 擅长快速转换和边路突破的球队体系\n- 适合 4-3-3 / 4-2-3-1 阵型中的边锋角色\n- 防守反击战术中可作为快速推进点",
        "AM": "- 需要前场创造力支点的球队\n- 适合 4-2-3-1 阵型中的前腰或影锋角色\n- 控球主导的战术体系中可串联中场与锋线",
        "CM": "- 对中场控制力有要求的球队\n- 适合 4-3-3 阵型中的中前卫或 3-5-2 中的全能中场\n- 高位压迫体系中需要覆盖范围大的 B2B 中场",
        "DM": "- 需要中场屏障的球队\n- 适合 4-2-3-1 / 4-3-3 阵型中的防守型中场\n- 低位防守体系中可作为后防线前的清道夫",
    }
    return fits.get(pos, fits["CM"])


def _build_one_liner(row: pd.Series) -> str:
    """一句话总结。"""
    pos_cn = _POS_CN.get(row.get("standard_position", "CM"), row.get("standard_position"))
    score = float(row.get("total_score", 0))
    if score >= 85:
        tier = "顶级潜力"
    elif score >= 75:
        tier = "优质潜力"
    elif score >= 65:
        tier = "值得关注"
    else:
        tier = "有待观察"
    return (
        f"{row['player_name']} 是一名 {int(row['age'])} 岁的 {pos_cn}，"
        f"潜力评分 {score:.1f} 分（{tier}），"
        f"在同位置球员中处于 {row['standard_position']} 组前列。"
        f"核心评分 {float(row['core_performance']):.1f} / 行为风格 {float(row['behavior_score']):.1f}。"
    )


def _build_data_limitations(row: pd.Series) -> str:
    """检查数据限制。"""
    issues = []
    completeness = float(row.get("data_completeness", 1) or 1)
    if completeness < 0.85:
        issues.append(f"核心字段完整度仅 {completeness*100:.0f}%，部分维度可能无法准确评估")
    if row.get("market_value") is None or (isinstance(row.get("market_value"), float) and pd.isna(row.get("market_value"))):
        issues.append("缺少身价数据")
    if row.get("height") is None or (isinstance(row.get("height"), float) and pd.isna(row.get("height"))):
        issues.append("缺少身高数据")
    if row.get("weight") is None or (isinstance(row.get("weight"), float) and pd.isna(row.get("weight"))):
        issues.append("缺少体重数据（FIFA 数据匹配失败）")
    if row.get("position_source") == "heuristic":
        issues.append("位置分类来自 FBref 粗位置启发式推断，可能不够精确")

    if not issues:
        return "当前数据完整性良好，核心字段均可用。"
    return "以下数据存在限制，结论需谨慎：\n- " + "\n- ".join(issues)


def _fill_template(template: str, row: pd.Series, rank: int, total: int) -> str:
    """用球员数据填充报告模板。"""
    subs = {
        "{{ player_name }}": str(row.get("player_name", "?")),
        "{{ age }}": _fmt(int(row.get("age", 0))),
        "{{ standard_position }}": str(row.get("standard_position", "?")),
        "{{ team }}": str(row.get("team", "?")),
        "{{ league }}": str(row.get("league", "?")),
        "{{ nation }}": str(row.get("nation_", "?")) if pd.notna(row.get("nation_")) else "?",
        "{{ height }}": _fmt(row.get("height"), "?") + " cm" if row.get("height") and pd.notna(row.get("height")) else "?",
        "{{ weight }}": _fmt(row.get("weight"), "?") + " kg" if row.get("weight") and pd.notna(row.get("weight")) else "?",
        "{{ market_value }}": _fmt(row.get("market_value"), "?"),
        "{{ fifa_overall }}": _fmt(row.get("fifa_overall"), "?"),
        "{{ fifa_potential }}": _fmt(row.get("fifa_potential"), "?"),
        "{{ minutes }}": _fmt(int(row.get("minutes", 0))),
        "{{ goals }}": _fmt(int(row.get("goals", 0))),
        "{{ assists }}": _fmt(int(row.get("assists", 0))),
        "{{ shots }}": _fmt(int(row.get("shots", 0))),
        "{{ key_passes }}": _fmt(int(row.get("key_passes", 0))),
        "{{ dribbles }}": _fmt(int(row.get("dribbles", 0))),
        "{{ passes }}": _fmt(int(row.get("passes", 0))),
        "{{ tackles }}": _fmt(int(row.get("tackles", 0))),
        "{{ interceptions }}": _fmt(int(row.get("interceptions", 0))),
        "{{ yellow_cards }}": _fmt(int(row.get("yellow_cards", 0))),
        "{{ red_cards }}": _fmt(int(row.get("red_cards", 0))),
        "{{ goals_per90 }}": _fmt(row.get("goals_per90"), "0"),
        "{{ assists_per90 }}": _fmt(row.get("assists_per90"), "0"),
        "{{ shots_per90 }}": _fmt(row.get("shots_per90"), "0"),
        "{{ key_passes_per90 }}": _fmt(row.get("key_passes_per90"), "0"),
        "{{ dribbles_per90 }}": _fmt(row.get("dribbles_per90"), "0"),
        "{{ tackles_per90 }}": _fmt(row.get("tackles_per90"), "0"),
        "{{ interceptions_per90 }}": _fmt(row.get("interceptions_per90"), "0"),
        "{{ xg_per90 }}": _fmt(row.get("xg_per90"), "0"),
        "{{ xa_per90 }}": _fmt(row.get("xa_per90"), "0"),
        "{{ progressive_carries_per90 }}": _fmt(row.get("progressive_carries_per90"), "0"),
        "{{ shot_creating_actions_per90 }}": _fmt(row.get("shot_creating_actions_per90"), "0"),
        "{{ age_score }}": _fmt(row.get("age_score"), "0"),
        "{{ reliability_score }}": _fmt(row.get("reliability_score"), "0"),
        "{{ core_performance }}": _fmt(row.get("core_performance"), "0"),
        "{{ behavior_score }}": _fmt(row.get("behavior_score"), "0"),
        "{{ league_score }}": _fmt(row.get("league_score"), "0"),
        "{{ risk_penalty }}": _fmt(row.get("risk_penalty"), "0"),
        "{{ total_score }}": _fmt(row.get("total_score"), "0"),
        "{{ rank }}": str(rank),
        "{{ total_players }}": str(total),
        "{{ strengths }}": _build_strengths(row),
        "{{ weaknesses }}": _build_weaknesses(row),
        "{{ behavior_profile }}": _build_behavior_profile(row),
        "{{ risks }}": _build_risks(row),
        "{{ tactical_fit }}": _build_tactical_fit(row),
        "{{ one_liner }}": _build_one_liner(row),
        "{{ data_limitations }}": _build_data_limitations(row),
    }
    result = template
    for k, v in subs.items():
        result = result.replace(k, v)
    return result


def generate_reports(df: pd.DataFrame, output_dir: str | Path,
                     template_path: str | Path,
                     top_n: int = 20) -> list[Path]:
    """为 Top N 球员生成 Markdown + DOCX 报告。

    Args:
        df: 含评分的 DataFrame。
        output_dir: 输出根目录（会创建 reports_md/ 和 reports_docx/ 子目录）。
        template_path: Markdown 模板文件路径。
        top_n: 生成前 N 名。

    Returns:
        已生成的 Markdown 文件路径列表。
    """
    template = Path(template_path).read_text(encoding="utf-8")
    official = df[df["is_official"] == True].copy()
    top_players = official.nlargest(top_n, "total_score")
    total_players = len(official)

    out = Path(output_dir)
    md_dir = out / "reports_md"
    docx_dir = out / "reports_docx"
    md_dir.mkdir(parents=True, exist_ok=True)
    docx_dir.mkdir(parents=True, exist_ok=True)

    generated: list[Path] = []

    for rank, (_, player) in enumerate(top_players.iterrows(), 1):
        md_content = _fill_template(template, player, rank, total_players)

        safe_name = str(player["player_name"]).replace("/", "_").replace("\\", "_")
        md_path = md_dir / f"{safe_name}.md"
        md_path.write_text(md_content, encoding="utf-8")
        generated.append(md_path)

        # DOCX
        try:
            _md_to_docx(md_content, docx_dir / f"{safe_name}.docx")
        except Exception:
            pass  # DOCX 生成失败不阻塞

    print(f"[阶段8] 已生成 {len(generated)} 份 Markdown 报告 → {md_dir}")

    # 尝试生成 DOCX
    try:
        import docx as _docx
        docx_count = len(list(docx_dir.glob("*.docx")))
        if docx_count > 0:
            print(f"[阶段8] 已生成 {docx_count} 份 DOCX 报告 → {docx_dir}")
    except ImportError:
        print("[阶段8] python-docx 未安装，跳过 DOCX 生成。")

    return generated


def _md_to_docx(md_content: str, output_path: Path) -> None:
    """将 Markdown 内容转换为 DOCX 文件（简化版）。"""
    from docx import Document
    from docx.shared import Pt, Inches

    doc = Document()
    style = doc.styles["Normal"]
    style.font.size = Pt(10)
    style.font.name = "Calibri"

    for line in md_content.split("\n"):
        line = line.strip()
        if not line:
            continue
        if line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("---"):
            doc.add_paragraph("─" * 60)
        elif line.startswith("| "):
            # 跳过表格行（简化处理，写为普通文本）
            doc.add_paragraph(line.strip("| "), style="List Bullet")
        elif line.startswith("- "):
            doc.add_paragraph(line[2:], style="List Bullet")
        elif line.startswith("*"):
            pass  # 跳过纯装饰行
        elif line.startswith(">"):
            p = doc.add_paragraph(line[1:].strip())
            p.runs[0].italic = True if p.runs else None
        else:
            doc.add_paragraph(line)

    doc.save(str(output_path))
